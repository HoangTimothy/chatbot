from dataclasses import dataclass
from pathlib import Path
from typing import Protocol


@dataclass(frozen=True)
class LoadedDocument:
    document_id: str
    source_uri: str
    title: str
    text: str
    metadata: dict[str, str]


class DocumentLoader(Protocol):
    def load(self, path: Path) -> LoadedDocument:
        """Load a document into normalized text."""
        ...


class PdfDocumentLoader:
    def load(self, path: Path) -> LoadedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        return LoadedDocument(
            document_id=path.stem,
            source_uri=str(path),
            title=path.stem,
            text=text,
            metadata={"file_type": "pdf"},
        )


class DocxDocumentLoader:
    def load(self, path: Path) -> LoadedDocument:
        from docx import Document

        doc = Document(str(path))
        text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return LoadedDocument(
            document_id=path.stem,
            source_uri=str(path),
            title=path.stem,
            text=text,
            metadata={"file_type": "docx"},
        )


class LoaderRegistry:
    def __init__(self) -> None:
        self._loaders: dict[str, DocumentLoader] = {
            ".pdf": PdfDocumentLoader(),
            ".docx": DocxDocumentLoader(),
        }

    def load(self, path: Path) -> LoadedDocument:
        loader = self._loaders.get(path.suffix.lower())
        if loader is None:
            raise ValueError(f"Unsupported document type: {path.suffix}")
        return loader.load(path)
