import csv
from typing import Sequence
import pathlib
import pypdf
import docx
import openpyxl
from bs4 import BeautifulSoup

from rag_core.contracts.types import ParsedBlock
from rag_core.ports.interfaces import DocumentParserPort


class TextParser(DocumentParserPort):
    """Parses plain text files into paragraphs."""

    def parse(self, file_path: str) -> Sequence[ParsedBlock]:
        path = pathlib.Path(file_path)
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        # Split by double newlines into blocks, stripping extra spaces
        raw_blocks = content.split("\n\n")
        blocks = []
        for block in raw_blocks:
            cleaned = block.strip()
            if cleaned:
                blocks.append(ParsedBlock(text=cleaned, page_number=1))
        return blocks


class CsvParser(DocumentParserPort):
    """Parses CSV tabular files into structured text blocks per row."""

    def parse(self, file_path: str) -> Sequence[ParsedBlock]:
        path = pathlib.Path(file_path)
        blocks = []
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            headers = next(reader, None)
            if not headers:
                return []

            # Clean headers
            headers = [h.strip() for h in headers]

            for row_idx, row in enumerate(reader):
                cleaned_row = [val.strip() for val in row]
                # Format: "Column1: value1 | Column2: value2"
                row_str = " | ".join(
                    f"{headers[i]}: {cleaned_row[i]}"
                    for i in range(min(len(headers), len(cleaned_row)))
                )
                if row_str.strip():
                    blocks.append(
                        ParsedBlock(
                            text=row_str,
                            page_number=1,
                            section_title="Data Table",
                        )
                    )
        return blocks


class PDFParser(DocumentParserPort):
    """Parses PDF documents, extracting text per page."""

    def parse(self, file_path: str) -> Sequence[ParsedBlock]:
        path = pathlib.Path(file_path)
        blocks = []
        reader = pypdf.PdfReader(str(path))

        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            cleaned_text = text.strip() if text else ""
            # Note: We do not fail on empty pages, but record them for downstream warnings
            blocks.append(ParsedBlock(text=cleaned_text, page_number=idx + 1))

        return blocks


class DocxParser(DocumentParserPort):
    """Parses Word documents, extracting heading hierarchy and paragraphs."""

    def parse(self, file_path: str) -> Sequence[ParsedBlock]:
        path = pathlib.Path(file_path)
        doc = docx.Document(str(path))
        blocks = []

        current_heading_path = []
        current_section = None

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower()

            if "heading" in style_name:
                # Try to extract level, e.g. "heading 1" -> 1
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                # Clean heading path up to current level
                current_heading_path = current_heading_path[: level - 1]
                current_heading_path.append(text)
                current_section = text
            else:
                blocks.append(
                    ParsedBlock(
                        text=text,
                        page_number=1,
                        section_title=current_section,
                        heading_path=tuple(current_heading_path),
                    )
                )

        # Handle tables in docx
        for table_idx, table in enumerate(doc.tables):
            table_text_rows = []
            for row in table.rows:
                row_vals = [cell.text.strip() for cell in row.cells]
                table_text_rows.append(" | ".join(row_vals))

            table_text = "\n".join(table_text_rows)
            if table_text.strip():
                blocks.append(
                    ParsedBlock(
                        text=table_text,
                        page_number=1,
                        section_title=f"Table {table_idx + 1}",
                        heading_path=tuple(current_heading_path),
                    )
                )

        return blocks


class XlsxParser(DocumentParserPort):
    """Parses Excel spreadsheets sheet-by-sheet."""

    def parse(self, file_path: str) -> Sequence[ParsedBlock]:
        path = pathlib.Path(file_path)
        wb = openpyxl.load_workbook(str(path), data_only=True)
        blocks = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_rows = []
            for row in sheet.iter_rows(values_only=True):
                if any(val is not None for val in row):
                    row_str = " | ".join(
                        str(val).strip() if val is not None else "" for val in row
                    )
                    sheet_rows.append(row_str)

            sheet_text = "\n".join(sheet_rows)
            if sheet_text.strip():
                blocks.append(
                    ParsedBlock(
                        text=sheet_text,
                        page_number=1,
                        section_title=f"Sheet: {sheet_name}",
                    )
                )

        return blocks


class HtmlParser(DocumentParserPort):
    """Parses HTML files, stripping styling/scripts and extracting body paragraphs."""

    def parse(self, file_path: str) -> Sequence[ParsedBlock]:
        path = pathlib.Path(file_path)
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        # Strip scripts and styles
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()

        blocks = []
        current_heading_path = []
        current_section = None

        # Iterate structural tags linearly
        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
            text = element.get_text().strip()
            if not text:
                continue

            tag_name = element.name
            if tag_name.startswith("h"):
                try:
                    level = int(tag_name[1])
                except (ValueError, IndexError):
                    level = 1

                current_heading_path = current_heading_path[: level - 1]
                current_heading_path.append(text)
                current_section = text
            else:
                blocks.append(
                    ParsedBlock(
                        text=text,
                        page_number=1,
                        section_title=current_section,
                        heading_path=tuple(current_heading_path),
                    )
                )

        return blocks


class ParserRegistry:
    """Registry to select and cache parser instances by file extension."""

    def __init__(self):
        self._parsers = {
            ".txt": TextParser(),
            ".csv": CsvParser(),
            ".pdf": PDFParser(),
            ".docx": DocxParser(),
            ".xlsx": XlsxParser(),
            ".html": HtmlParser(),
            ".htm": HtmlParser(),
        }

    def get_parser(self, file_extension: str) -> DocumentParserPort:
        ext = file_extension.lower()
        if not ext.startswith("."):
            ext = f".{ext}"
        parser = self._parsers.get(ext)
        if not parser:
            raise ValueError(f"Unsupported file type extension: '{file_extension}'")
        return parser
